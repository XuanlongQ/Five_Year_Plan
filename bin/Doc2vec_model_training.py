# -*- encoding: utf-8 -*-
'''
@file: Doc2vec_model_training.py
@Author: Xuanlong
@email: qxlpku@gmail.com
''' 

# This file is used to train the Doc2Vec model on the 12th Five Year Plan documents

from gensim.models.doc2vec import Doc2Vec, TaggedDocument
from docx import Document
import jieba
import re
import os
import pandas as pd

# configuration - stop words
stopWordsPath = "Five_Year_Plan/bin/conf/stopwords"

class Docx:
    def __init__(self, input_file, load_saved=False):
        self.input_file = input_file
        self.name = self.get_name(input_file)
        self.file = Document(input_file)
        if not load_saved:
            self.doc_raw_text = self.get_doc_raw_text()
            self.doc_token_text = self.get_doc_token_text()
        else:
            self.doc_raw_text = ""
            self.doc_token_text = []

    def get_name(self, input_file):
        _, file_name = os.path.split(input_file)
        name, _ = os.path.splitext(file_name)
        return name

    def get_doc_raw_text(self):
        tmp = []
        for para in self.file.paragraphs:
            tmp.append(para.text.strip().replace(" ", "").replace("\t", ""))
        return "".join(tmp)

    def _get_text_token(self, text):
        token_list = list(jieba.cut(re.sub('[^\u4e00-\u9fa5]+', '', text), use_paddle=False))
        res = []
        for token in token_list:
            if token in STOPWORDS:
                continue
            res.append(token)
        return res

    def get_doc_token_text(self):
        return self._get_text_token(self.doc_raw_text)

def get_stopwords(filefolder):
    stopwords_list = []
    for file in os.listdir(filefolder):
        stopwords_file = os.path.join(filefolder, file)
        with open(stopwords_file,"r",encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                stopwords_list.append(line)
    return stopwords_list

def getFiles(input_directory):
    file_names = sorted([f for f in os.listdir(input_directory) if f.endswith('.docx')])
    files = []
    for filename in file_names:
        # if filename.endswith(".docx"):
        files.append(os.path.join(input_directory, filename))
    return files

def append_to_dataframe(tag,document_vec): 
    
    pklPath = 'Five_Year_Plan/bin/doc/Doc2vec_Models/tag_document_vecs_100v.pkl'
    try:
        # Load the existing DataFrame
        df = pd.read_pickle(pklPath)
    except FileNotFoundError:
        df = pd.DataFrame(columns=['cities','dimension'])

    # Create a new DataFrame with the data to be appended
    new_row = pd.DataFrame([[tag,document_vec]], columns=['cities','dimension'])
    
    # Concatenate the new DataFrame with the existing DataFrame
    df = pd.concat([df, new_row], ignore_index=True)
    
    # Save the updated DataFrame
    df.to_pickle(pklPath)

if __name__ == "__main__":
    # the input directory, you could get the docx files from the following link,
    input_directory = "Five_Year_Plan/bin/doc/12th_fyp_all"
    
    # Models will be saved in the following directory
    Doc2Vec_ModelPath = "Five_Year_Plan/bin/doc/Doc2vec_Models/5yrPlans_model_100v.bin"
    
    STOPWORDS = get_stopwords(stopWordsPath)
    input_files = getFiles(input_directory)
    print("{} files were added".format(len(input_files)))

    ## Preparing corpus
    corpus = []
    corpus_names = []
    files_count = 0
    for input_file in input_files:
        print("Processing: {}".format(input_file))
        doc = Docx(input_file)
        files_count += 1
        corpus.append(doc.doc_token_text)
        corpus_names.append(doc.name)

    print("Total files processed: {}".format(files_count))
    ## Model Training
    tagged_data = [TaggedDocument(words=doc, tags=[name]) for doc, name in zip(corpus, corpus_names)]
    
    model = Doc2Vec(vector_size=100, window=5, min_count=3, workers=1, epochs=20, seed = 42)
    model.build_vocab(tagged_data)
    model.train(tagged_data, total_examples=model.corpus_count,epochs=20)
    
    for tag in corpus_names:
        vector = model.dv[tag]
        append_to_dataframe(tag,vector)

    # Save the trained model
    model.save(Doc2Vec_ModelPath)
    

    

    
    

